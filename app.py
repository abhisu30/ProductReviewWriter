import streamlit as st
import pandas as pd
import TokenCounter as TC
from openai import OpenAI
from docx import Document
import time
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import os
from prompts import prompts_list

#Function to create Sectional summary

def GPTSummarizer(sec, rev, api_key, model):

    client = OpenAI(
    api_key=api_key,
    )
    MODEL = "gpt-3.5-turbo-1106"

    #print("Summarizing section  = " + sec)

    instruction = prompts_list["summarizePrompt"].format(sec=sec, rev=rev)
    input = str(instruction)
    messages=[
        {"role": "system", "content": "You are an expert summarizer."},
        {"role": "user", "content": input},
    ]
    response = client.chat.completions.create(
    model=MODEL,
    messages=messages,
    temperature=0.2,
    )

    rev_output = response.choices[0].message.content
    response_token = response.usage.prompt_tokens
    token_count = TC.num_tokens_from_messages(messages, MODEL) + response_token

    #print('\nSummary inside function = ' + rev_output)
    
    return rev_output, token_count

#Function to create main content

def GPTReviewCreator(cat, revSum, pname, pkw, KList, sec, rev, prev_para, nxtSec, api_key, model):
    messages =[]
    kwCluster = ', '.join(KList)
    client = OpenAI(
    api_key=api_key,
    )
    MODEL = model
    limited_text = str(prev_para[:200]) if prev_para is not None else ''
    print("Previous Paragraph = " + str(limited_text))

    instruction = prompts_list["reviewPrompt"].format(pname=pname, cat=cat, sec=sec, nxtSec=nxtSec, pkw=pkw, kwCluster=kwCluster, revSum=revSum, rev=rev)
    input = str(instruction)
    messages=[
        {"role": "system", "content": "You are a trusted product reviewer."},
        {"role": "user", "content": input},
    ]
    response = client.chat.completions.create(
    model=MODEL,
    messages=messages,
    temperature=0.4,
    )

    rev_output = response.choices[0].message.content
    response_token = response.usage.prompt_tokens
    token_count = TC.num_tokens_from_messages(messages, MODEL) + response_token
    
    return rev_output, token_count

#Function to create Introduction

def GPTReviewIntroCreator(cat, pname, pkw, KList, sec, rev, prev_para, nxtSec, api_key, model):
    messages =[]
    kwCluster = ', '.join(KList)
    client = OpenAI(
    api_key=api_key,
    )
    MODEL = model
    instruction = prompts_list["introPrompt"].format(pname=pname, cat=cat, sec=sec, nxtSec=nxtSec, pkw=pkw, rev=rev, prev_para=prev_para)
    input = str(instruction)
    messages=[
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": input},
    ]
    response = client.chat.completions.create(
    model=MODEL,
    messages=messages,
    temperature=0.6,
    )

    rev_output = response.choices[0].message.content
    response_token = response.usage.prompt_tokens
    token_count = TC.num_tokens_from_messages(messages, MODEL) + response_token
    
    return rev_output, token_count

def Points2Reviews(PCat, PName, Pkw, KList, workbook, api_key, model):
    token_count = 0
    document = Document()

    # Initialize the DataFrame for Excel
    df = pd.DataFrame(columns=["Section", "Review"])

    # Assuming the first sheet contains the data
    first_sheet_name = list(workbook.keys())[0]
    worksheet_review = workbook[first_sheet_name]

    prev_para = 'None'
    nxtSec = 'None'
    revSum = ''
    for i, row in worksheet_review.iterrows():
        sec = str(row['Section'])
        rev = str(row['Review'])

        if not sec or not rev:
            break

        # To get the next and previous rows, use the DataFrame directly
        if i < len(worksheet_review) - 1:
            nxtSec = str(worksheet_review.iloc[i + 1]['Section']) if worksheet_review.iloc[i + 1]['Section'] else None
        else:
            nxtSec = None

        prev_para = str(worksheet_review.iloc[i - 1]['Section']) if i > 0 else 'None'

        if sec == 'None' or sec == '':
            break

        print('\nCurrent Section Name: ' + str(sec))
        print('Next Section Name: ' + str(nxtSec))

        if str(sec).lower() == 'introduction':
        # Call the GPTReviewIntroCreator function to create a review Introduction
            rev_op, OPtokens = GPTReviewIntroCreator(PCat, PName, Pkw, KList, sec, rev, prev_para, nxtSec, api_key, model)
            rev_op_summ, summary_token = GPTSummarizer(sec, rev, api_key, model)
            revSum = revSum + '\n' + str(rev_op_summ)
            print('\nIntro Summary: ' + revSum + '\n')
        else:
        # Call the GPTReviewCreator function to create a review
            rev_op, OPtokens = GPTReviewCreator(PCat, revSum, PName, Pkw, KList, sec, rev, prev_para, nxtSec, api_key, model)
            rev_op_summ, summary_token = GPTSummarizer(sec, rev, api_key, model)
            revSum = revSum + '\n' + str(rev_op_summ)
            print('\nSummary: ' + revSum + '\n')

        # Write the section title in bold text
        document.add_paragraph(sec, style='Heading 1')
        
        # Write the output text under the section title
        paragraph = document.add_paragraph(rev_op)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(10)
        paragraph_format.space_after = Pt(10)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Save the new header and content to the DataFrame
        new_row = pd.DataFrame({"Section": [sec], "Review": [rev_op]})
        df = pd.concat([df, new_row], ignore_index=True)

        if str(sec).lower() == 'introduction':
            token_count = token_count + OPtokens
        else:
            token_count = token_count + OPtokens + summary_token

        # introduce a delay of 20 seconds between iterations
        time.sleep(18)  
        #print('Section '+ str(sec) + ' completed\n') 

         # Return the document and the token count

    output_folder = 'Output'
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    
    # Save the Excel file
    excel_filename = os.path.join(output_folder, "reviews.xlsx")
    df.to_excel(excel_filename, index=False)

    # Save the Word file
    doc_filename = os.path.join(output_folder, "reviews_doc.docx")
    document.save(doc_filename)


    print(f"Tokens spent: {token_count}")
    result = token_count * 0.002 / 1000
    print(f"Total cost: {result}")

    return doc_filename, excel_filename, token_count

def main():
    st.title("Product Review Generator")

    # Streamlit widgets to take user inputs
    model_selection = st.selectbox("Select the Model*", ["GPT 3.5", "GPT-4"], index=1)
    api_key = st.text_input("API Key*", type="password")
    pname = st.text_input("Product Name*")
    pcat = st.text_input("Product Category*")
    pkw = st.text_input("Primary Keyword*")
    seckw = st.text_input("Secondary Keywords (separated by commas)")

    # Model mapping
    model_mapping = {
        "GPT 3.5": "gpt-3.5-turbo-1106",
        "GPT-4": "gpt-4-0613"
    }
    model = model_mapping[model_selection]

    def get_base64_encoded_data(filename):
        with open(filename, "rb") as file:
            return base64.b64encode(file.read()).decode()

    # Provide a hyperlink for the template
    template_file = 'Product_Review_Input_Template.xlsx'
    if os.path.exists(template_file):
        base64_data = get_base64_encoded_data(template_file)
        href = f'<a href="data:application/octet-stream;base64,{base64_data}" download="{template_file}">Download Input Template</a>'
        st.markdown(href, unsafe_allow_html=True)

    # File uploader for excel file
    uploaded_file = st.file_uploader("Upload filled Excel sheet", type="xlsx")

    # Process the inputs if the user clicks the 'Generate Review' button
    if st.button("Generate Review"):
        if api_key and pname and pcat and pkw and uploaded_file:
            # Load the Excel file
            workbook = pd.read_excel(uploaded_file, sheet_name=None)
            
            # Split the secondary keywords
            KList = seckw.split(',') if seckw else []

            # Call the Points2Reviews function
            doc_filename, excel_filename, token_count = Points2Reviews(pcat, pname, pkw, KList, workbook, api_key, model)

            st.success("Content Generated Successfully!")
            st.write(f"Tokens spent: {token_count}")
            result = token_count * 0.002 / 1000
            st.write(f"Total cost: {result}")

            # Provide download links for generated files
            with open(doc_filename, "rb") as file:
                btn = st.download_button(
                    label="Download Review Document",
                    data=file,
                    file_name=doc_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            with open(excel_filename, "rb") as file:
                btn = st.download_button(
                    label="Download Excel File",
                    data=file,
                    file_name=excel_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
        else:
            st.error("Please fill in all required fields and upload the Excel file.")

if __name__ == "__main__":
    main()